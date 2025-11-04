# app/routes.py
import os
import re
import io
import uuid
import io
import google.cloud.storage
from datetime import timedelta
from flask import (
    render_template, request, redirect, url_for, session, 
    send_file, jsonify
)
from docx import Document
from werkzeug.utils import secure_filename
from app import app  # Import the 'app' object from __init__.py
from app.helpers import generate_smart_question # Import our helper

storage_client = google.cloud.storage.Client()

# --- Homepage Route ---
@app.route('/')
def index():
    return render_template('index.html')

# --- API route for fast upload and preview ---
@app.route('/api/upload-preview', methods=['POST'])
def api_upload_preview():
    if 'document' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['document']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file:
        try:
            filename = secure_filename(file.filename)
            filepath = f"uploads/{filename}" 
            
            # --- Read the file into memory ONCE ---
            file_bytes = file.read()
            # ----------------------------------------------

            # 1. Use the bytes to upload to GCS
            bucket = storage_client.bucket(app.config['GCS_BUCKET_NAME'])
            blob = bucket.blob(filepath)
            # Create a new in-memory stream for GCS
            gcs_stream = io.BytesIO(file_bytes)
            blob.upload_from_file(gcs_stream)

            # Save the GCS path to the session
            session['filepath'] = filepath 
            
            # 2. Use the bytes again to create the preview
            # Create a second in-memory stream for python-docx
            docx_stream = io.BytesIO(file_bytes)
            document = Document(docx_stream)
            # ----------------------------------------------
            
            preview_text_list = [para.text for para in document.paragraphs]
            full_preview = "\n".join(preview_text_list)
            
            return jsonify({'preview_text': full_preview})
        
        except Exception as e:
            print(f"Error parsing GCS preview: {e}")
            return jsonify({'error': 'Could not parse .docx file.'}), 500
    
    return jsonify({'error': 'File upload failed.'}), 500

# --- API route for slow AI processing ---
@app.route('/api/process-ai', methods=['GET'])
def api_process_ai():
    filepath = session.get('filepath')
    if not filepath:
        return jsonify({'error': 'No file in session.'}), 400
    
    try:
        # Get the file from GCS
        bucket = storage_client.bucket(app.config['GCS_BUCKET_NAME'])
        blob = bucket.blob(filepath)
        file_bytes = blob.download_as_bytes()
        file_stream = io.BytesIO(file_bytes)
        
        document = Document(file_stream)
        
        # --- PARSING LOGIC ---
        
        pattern = r'\[.*?\]' 
        questions = []
        question_id = 0

        all_paragraphs = [para.text for para in document.paragraphs]
        full_document_text = "\n".join(all_paragraphs)
        
        print("--- Starting AI question generation in background... ---")
        
        # Helper function to process placeholders
        def process_placeholders(para):
            nonlocal question_id
            found_in_para = re.findall(pattern, para.text)
            for p in found_in_para:
                context = para.text
                ai_question, ai_explanation = generate_smart_question(
                    full_document_text, context, p
                )
                question = {
                    'id': question_id,
                    'placeholder': p,
                    'context': context,
                    'ai_question': ai_question,
                    'ai_explanation': ai_explanation
                }
                questions.append(question)
                question_id += 1

        # Loop 1: Find in all paragraphs
        for para in document.paragraphs:
            process_placeholders(para)
        
        # Loop 2: Find in all tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        process_placeholders(para)

        print(f"--- Finished. Found {len(questions)} questions. ---")

        # 7. Save the questions list
        session['questions'] = questions
        session['answers'] = {} 
        
        return jsonify({'status': 'ready'})

    except Exception as e:
        print(f"Error in api_process_ai: {e}")
        return jsonify({'error': 'Failed during AI processing.'}), 500

# --- Form Filling (Chat) Page ---
@app.route('/fill')
def fill_form():
    filepath = session.get('filepath')
    if not filepath:
        return redirect(url_for('index'))
        
    try:
        # Get the file from GCS
        bucket = storage_client.bucket(app.config['GCS_BUCKET_NAME'])
        blob = bucket.blob(filepath)
        file_bytes = blob.download_as_bytes()
        file_stream = io.BytesIO(file_bytes)

        document = Document(file_stream)
        preview_text_list = [para.text for para in document.paragraphs]
        preview_text = "\n".join(preview_text_list)
    except Exception as e:
        print(f"Error loading preview for chat: {e}")
        preview_text = "Error: Could not load document preview."
    
    return render_template('chat.html', preview_text=preview_text)

# --- API Route for the Chat Interface ---
@app.route('/api/chat', methods=['GET', 'POST'])
def api_chat():
    questions = session.get('questions', [])
    answers = session.get('answers', {})

    if request.method == 'POST':
        data = request.json
        answer = data.get('answer')
        question_id = data.get('id')
        
        if question_id is not None:
            answers[str(question_id)] = answer
            session['answers'] = answers
    
    next_question = None
    for q in questions:
        if str(q['id']) not in answers:
            next_question = q
            break
            
    if next_question:
        return jsonify({
            'status': 'question',
            'question_text': next_question['ai_question'],
            'explanation_text': next_question['ai_explanation'],
            'placeholder': next_question['placeholder'],
            'id': next_question['id'],
            'context': next_question['context']
        })
    else:
        return jsonify({
            'status': 'done',
            'redirect_url': url_for('generate_document')
        })

# --- Document Generation and Download Page ---
@app.route('/download')
def generate_document():
    try:
        filepath = session.get('filepath') # e.g., "uploads/mydoc.docx"
        questions = session.get('questions', [])
        answers = session.get('answers', {})

        if not filepath or not questions or not answers:
            return redirect(url_for('index'))

        # Download original file from GCS
        bucket = storage_client.bucket(app.config['GCS_BUCKET_NAME'])
        original_blob = bucket.blob(filepath)
        file_bytes = original_blob.download_as_bytes()
        file_stream = io.BytesIO(file_bytes)
        
        document = Document(file_stream)
        
        question_index = 0
        total_questions = len(questions)

        # Loop 1: Replace in all paragraphs
        for para in document.paragraphs:
            if question_index >= total_questions: break
            
            # Get the current question
            current_q = questions[question_index]
            
            # Keep checking the same paragraph until no more of its placeholders are found
            while current_q['placeholder'] in para.text:
                answer = answers.get(str(current_q['id']), '')
                
                # Use .replace() with count=1 to only replace the *first* instance
                para.text = para.text.replace(current_q['placeholder'], answer, 1)
                
                question_index += 1 # Move to the next question
                if question_index >= total_questions: break
                current_q = questions[question_index] # Get the next question

        # Loop 2: Replace in all tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if question_index >= total_questions: break
                        current_q = questions[question_index]
                        while current_q['placeholder'] in para.text:
                            answer = answers.get(str(current_q['id']), '')
                            para.text = para.text.replace(current_q['placeholder'], answer, 1)
                            
                            question_index += 1
                            if question_index >= total_questions: break
                            current_q = questions[question_index]
                        if question_index >= total_questions: break
                    if question_index >= total_questions: break
                if question_index >= total_questions: break
            if question_index >= total_questions: break

        # --- New Save Logic ---
        final_filename = f"completed_{uuid.uuid4()}.docx"
        final_filepath_in_gcs = f"completed/{final_filename}"
        
        # Save the edited document to a memory stream
        completed_stream = io.BytesIO()
        document.save(completed_stream)
        completed_stream.seek(0)
        
        # Upload the completed file to GCS
        completed_blob = bucket.blob(final_filepath_in_gcs)
        completed_blob.upload_from_file(completed_stream)
        
        # --- Preview Logic ---
        completed_stream.seek(0)
        preview_doc = Document(completed_stream)
        preview_text_list = [para.text for para in preview_doc.paragraphs]
        full_preview = "\n".join(preview_text_list)

        return render_template(
            'download.html', 
            preview_text=full_preview, 
            final_filename=final_filepath_in_gcs # Pass the GCS path
        )
    except Exception as e:
        print(f"Error generating document: {e}")
        return "An error occurred while generating your document."

# --- Serves the generated file for download ---
@app.route('/get-file/<path:filename>') # Use <path:> to capture slashes
def get_file(filename):
    try:
        bucket = storage_client.bucket(app.config['GCS_BUCKET_NAME'])
        blob = bucket.blob(filename)

        if not blob.exists():
            print(f"Error: Blob does not exist at path: {filename}")
            return "Error: File not found.", 404

        # Get the app's own service account email from its config
        # This is the identity you gave the 'Token Creator' role
        service_account_email = app.config.get('SERVICE_ACCOUNT_EMAIL')
        if not service_account_email:
             # This is your default service account from the error log
             service_account_email = "213433359152-compute@developer.gserviceaccount.com"

        signed_url = blob.generate_signed_url(
            version="v4",
            expiration=timedelta(minutes=15),
            method="GET",
            response_disposition="attachment; filename=completed_document.docx",
            service_account_email=service_account_email # <-- Add this line
        )
        
        return redirect(signed_url)
        
    except Exception as e:
        print(f"CRITICAL Error generating signed URL: {e}")
        return "An error occurred while trying to generate your download link.", 500



# --- Lets the user edit their answers ---
@app.route('/edit')
def edit_answers():
    session.pop('answers', None)
    return redirect(url_for('fill_form'))
